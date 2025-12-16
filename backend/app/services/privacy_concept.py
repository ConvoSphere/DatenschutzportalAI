import os
import logging
import json
import pypdf
import docx
import openpyxl
from typing import List, Optional
from datetime import datetime

from pydantic_ai import Agent
from app.config import settings
from app.models.privacy_concept import ExtractedStudyData

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from app.models.db_models import PrivacyConceptDB

logger = logging.getLogger(__name__)

class PrivacyConceptService:
    def __init__(self, db: Optional[AsyncSession] = None):
        self.db = db
        # Configure OpenAI env vars for Pydantic AI
        if settings.ai_api_key:
            os.environ["OPENAI_API_KEY"] = settings.ai_api_key
        if settings.ai_api_base_url:
            os.environ["OPENAI_BASE_URL"] = settings.ai_api_base_url
        if settings.ai_proxy:
            os.environ["HTTP_PROXY"] = settings.ai_proxy
            os.environ["HTTPS_PROXY"] = settings.ai_proxy

        self.extraction_agent = Agent(
            model=settings.ai_model_name,
            system_prompt="""Du bist ein Datenschutzexperte für medizinische Forschung an der Universitätsmedizin Frankfurt (UMF).
Analysiere den vorliegenden Forschungsantrag präzise und extrahiere die für das Datenschutzkonzept relevanten Metadaten.

WICHTIGE HINWEISE ZUR EXTRAKTION:
- Studientyp: Unterscheide genau zwischen 'retrospektiv' (nur Bestandsdaten), 'prospektiv' (neue Datenerhebung) oder 'gemischt'.
- Datenquellen: Achte auf Begriffe wie 'Orbis', 'iBDF', 'Klinisches Arbeitsplatzsystem', 'Patientenakte'.
- Pseudonymisierung: Suche nach Hinweisen auf 'Treuhandstelle', 'ID-Liste', 'Code-Key'.
- Institution: Falls nicht anders genannt, gehe von 'Universitätsmedizin Frankfurt' aus.

Antworte AUSSCHLIESSLICH mit dem geforderten JSON-Objekt.""",
            result_type=ExtractedStudyData,
        )

        self.generation_agent = Agent(
            model=settings.ai_model_name,
            system_prompt="""Du bist der Datenschutzbeauftragte der Universitätsmedizin Frankfurt (UMF).
Deine Aufgabe ist das Verfassen eines professionellen, behördenreifen Datenschutzkonzepts für einen Forschungsantrag.

STIL & TON:
- Formale, juristisch präzise Amtssprache (Deutsch).
- Sachlich, objektiv, direkt.
- Verwende die korrekten rechtlichen Bezüge: DSGVO (Datenschutz-Grundverordnung) und HDSIG (Hessisches Datenschutz- und Informationsfreiheitsgesetz).

FORMATIERUNG:
- Nutze Markdown (# Überschriften).
- Keine Platzhalter wie [Hier Datum einfügen] - fülle alles basierend auf den Daten oder sinnvollen Standards aus.""",
            result_type=str,
        )

    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text based on file extension."""
        ext = os.path.splitext(file_path)[1].lower()
        try:
            if ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif ext in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            else:
                logger.warning(f"Unsupported file type for text extraction: {ext}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""

    def _extract_from_pdf(self, path: str) -> str:
        text = ""
        with open(path, 'rb') as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        return text

    def _extract_from_docx(self, path: str) -> str:
        doc = docx.Document(path)
        return "\n".join([p.text for p in doc.paragraphs])

    async def extract_data(self, file_paths: List[str], manual_text: Optional[str] = None) -> ExtractedStudyData:
        combined_text = ""
        if manual_text:
            combined_text += f"\n\n--- MANUAL TEXT ---\n\n{manual_text}"
            
        for file_path in file_paths:
            text = self.extract_text_from_file(file_path)
            filename = os.path.basename(file_path)
            if text:
                combined_text += f"\n\n--- FILE: {filename} ---\n\n{text}"
        
        if not combined_text.strip():
            raise ValueError("No text provided for extraction.")

        prompt = f"""
        Analysiere den folgenden Forschungsantrag und extrahiere die relevanten Daten:
        
        {combined_text[:100000]}
        """
        
        result = await self.extraction_agent.run(prompt)
        return result.data

    async def generate_concept(self, data: ExtractedStudyData) -> str:
        prompt = f"""
        Erstelle ein detailliertes Datenschutzkonzept für folgende Studie:
        
        # STUDIENDATEN
        Titel: {data.study_title}
        Typ: {data.study_type}
        PI: {data.principal_investigator}
        Institution: {data.institution}
        Ziel: {data.study_goal}
        Datenarten: {', '.join(data.data_types)}
        Patientenzahl: {data.patient_count}
        Quellen: {', '.join(data.data_sources)}
        Verarbeitung: {data.processing_methods}
        Pseudonymisierung: {'Ja' if data.pseudonymization_usage else 'Nein'}
        Externe Weitergabe: {'Ja' if data.external_data_sharing else 'Nein'}
        Ethikvotum: {data.ethics_vote or 'Beantragt'}
        
        # ZUSATZINFOS
        Minimierung: {data.data_minimization or 'Es werden nur die für die Forschungsfrage unbedingt erforderlichen Daten erhoben (Grundsatz der Datenminimierung).'}
        Speicherort: {data.storage_location or 'U:\\Klifo (Geschütztes Netzlaufwerk der UMF)'}
        Archivierung: {data.archiving_period or '10 Jahre nach Abschluss der Studie gemäß guter wissenschaftlicher Praxis'}
        Interne Zugriffe: {', '.join(data.internal_access) if data.internal_access else 'Nur autorisierte Mitglieder der Forschungsgruppe'}
        Externe Partner: {data.external_partners or 'Keine'}

        # ANWEISUNG ZUR STRUKTUR (Bitte exakt einhalten)

        ## 1. DARSTELLUNG DES FORSCHUNGSVORHABENS
        - Beschreibe Ziel und Zweck basierend auf den Studiendaten.
        - Begründe die Erforderlichkeit der Datenverarbeitung.
        - Falls retrospektiv: Erkläre, warum eine Einwilligung unverhältnismäßig wäre (HDSIG § 24).
        - Falls prospektiv: Erwähne die schriftliche Einwilligung der Patienten.

        ## 2. ORGANISATORISCHE STRUKTUR
        - Verantwortliche Stelle: Universitätsklinikum Frankfurt, Theodor-Stern-Kai 7, 60590 Frankfurt am Main.
        - Institutsleitung: {data.principal_investigator}
        - Datenschutzbeauftragter: Datenschutzbeauftragter der UMF (datenschutz@kgu.de).

        ## 3. BESCHREIBUNG DER DATENVERARBEITUNG
        - Art der Daten: {', '.join(data.data_types)}
        - Kreis der Betroffenen: Patienten der {data.institution}
        - Datenherkunft: {', '.join(data.data_sources)}
        - Datenfluss: Beschreibe den Weg der Daten von der Quelle (z.B. Orbis) in die Forschungsdatenbank.
        - Pseudonymisierung: Beschreibe das Verfahren (Trennung von ID und medizinischen Daten).

        ## 4. RECHTSGRUNDLAGEN
        - Nenne DSGVO Art. 6 Abs. 1 lit. e (öffentliches Interesse) sowie Art. 9 Abs. 2 lit. j (Forschungszwecke).
        - Nenne HDSIG § 24 (Verarbeitung zu wissenschaftlichen Forschungszwecken).
        - Falls Einwilligung vorliegt: DSGVO Art. 6 Abs. 1 lit. a und Art. 9 Abs. 2 lit. a.

        ## 5. RECHTE DER BETROFFENEN
        - Liste auf: Auskunftsrecht, Berichtigungsrecht, Löschrecht, Einschränkung der Verarbeitung, Widerspruchsrecht.
        - Hinweis: Einschränkungen dieser Rechte sind gemäß HDSIG möglich, wenn sie den Forschungszweck unmöglich machen würden.

        ## 6. ORGANISATORISCHE MAßNAHMEN
        - Verpflichtung auf Datengeheimnis.
        - Zugriffskonzepte (Need-to-know-Prinzip).
        - Regelmäßige Schulungen der Mitarbeiter.

        ## 7. TECHNISCHE MAßNAHMEN (TOMs)
        - Speicherung auf gesicherten Servern der UMF (keine lokale Speicherung auf Laptops).
        - Zugriffsschutz durch Passwörter und Active Directory.
        - Automatische Backups durch das Zentrum für Informations- und Medizintechnik (ZIM).
        - Verschlüsselung bei etwaigem Datentransfer.

        Antworte NUR mit dem Markdown-Text. Beginne direkt mit der Überschrift "# Datenschutzkonzept".
        """
        
        result = await self.generation_agent.run(prompt)
        return result.data

    def export_to_docx(self, markdown_text: str, output_path: str):
        doc = docx.Document()
        for line in markdown_text.split('\n'):
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.strip() == '':
                continue
            else:
                doc.add_paragraph(line)
        doc.save(output_path)

    async def save_concept(self, extracted_data: ExtractedStudyData, markdown: str, session_id: Optional[str] = None) -> str:
        if not self.db:
             raise ValueError("DB Session not initialized")
        
        db_obj = PrivacyConceptDB(
            extracted_data=extracted_data.model_dump(),
            concept_markdown=markdown,
            session_id=session_id
        )
        self.db.add(db_obj)
        await self.db.commit()
        await self.db.refresh(db_obj)
        return db_obj.id

    async def get_concept(self, concept_id: str) -> Optional[PrivacyConceptDB]:
        if not self.db:
             raise ValueError("DB Session not initialized")
        
        result = await self.db.execute(select(PrivacyConceptDB).where(PrivacyConceptDB.id == concept_id))
        return result.scalars().first()
